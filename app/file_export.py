import base64
import io
from datetime import datetime
import os
from docx.enum.section import WD_SECTION_START

from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from openai import OpenAI
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image as RLImage,
)
from PIL import Image
from app.templates.layouts import get_layout

load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def _decode_base64_image(img_b64: str) -> io.BytesIO:
    return io.BytesIO(base64.b64decode(img_b64))


def _build_kpi_summary_rows(kpi_data: Dict[str, Any]) -> List[List[str]]:
    rows: List[List[str]] = [["Metric", "Value"]]

    disabilities = kpi_data.get("Percentage of Employees with Disabilities", {}) or {}
    rows.append(
        [
            "Employees with Disabilities (Overall %)",
            f"{disabilities.get('overall_percentage', 0)}%",
        ]
    )

    turnover = kpi_data.get("Employee Turnover Rate", {}) or {}
    rows.append(
        ["Employee Turnover Rate", f"{turnover.get('overall_turnover_rate', 0)}%"]
    )

    training = kpi_data.get("Average Training Hours per Employee", {}) or {}
    rows.append(
        [
            "Average Training Hours per Employee",
            f"{training.get('overall_average_hours', 0)}",
        ]
    )

    injury = kpi_data.get("Workplace Injury Rate", {}) or {}
    rows.append(
        [
            "Workplace Injury Rate (per employee)",
            f"{injury.get('overall_injury_rate', 0)}",
        ]
    )

    workforce_by_gender = kpi_data.get("Total Workforce by Gender", []) or []
    for item in workforce_by_gender:
        rows.append(
            [
                f"Workforce - {item.get('gender', 'Unknown')}",
                str(item.get("employee_count", 0)),
            ]
        )

    return rows


def generate_docx_report(
    *,
    company_id: int,
    year: int,
    company_name: Optional[str],
    kpi_data: Dict[str, Any],
    charts: Dict[str, Optional[str]],
    executive_summary: str,
    workforce_composition_and_diversity: str,
    working_conditions_and_equal_opportunity: str,
    training_and_development: str,
    turnover_and_retention: str,
    health_and_safety: str,
    outlook_and_next_steps: str,
    closing: str,
    layout_template: Optional[Dict[str, Any]] = None,
) -> Tuple[str, str]:

    layout = get_layout("default") if layout_template is None else layout_template

    filename = f"S1_Report_{company_id}_{year}.docx"
    doc = Document()

    # ---- COVER PAGE ----
    cover_config = layout.get("cover_page", {})
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(cover_config.get("title", "ESRS S1 Management Report"))
    run.bold = True
    run.font.size = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().add_run("").add_break()

    sub = doc.add_paragraph(f"{company_name or company_id}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(16)

    doc.add_paragraph().add_run("").add_break()
    sub2 = doc.add_paragraph(f"Reporting Year: {year}")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.runs[0].font.size = Pt(14)

    # Logo placeholder
    if cover_config.get("show_logo", True):
        img = Image.new("RGB", (1024, 512), color=(230, 230, 230))
        img_buf = io.BytesIO()
        img.save(img_buf, format="PNG")
        img_buf.seek(0)
        doc.add_picture(img_buf, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date if enabled
    if cover_config.get("show_date", True):
        date_p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Confidential notice if enabled
    if cover_config.get("show_confidential", True):
        conf = doc.add_paragraph("Confidential")
        conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf.runs[0].italic = True

    doc.add_page_break()

    # Summary of KPI metrics
    kpi_config = layout.get("kpi_summary", {})
    doc.add_heading(kpi_config.get("title", "KPI Summary"), level=1)
    rows = _build_kpi_summary_rows(kpi_data)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = rows[0][0]
    hdr_cells[1].text = rows[0][1]
    for r in rows[1:]:
        row_cells = table.add_row().cells
        row_cells[0].text = r[0]
        row_cells[1].text = r[1]

    doc.add_paragraph("")

    # KPI visualizations
    visuals_config = layout.get("visuals", {})
    doc.add_heading(visuals_config.get("title", "KPI Visualizations"), level=1)

    # Display charts in configured order
    chart_order = visuals_config.get(
        "chart_order",
        [
            "workforce_by_gender",
            "training_hours_by_gender",
            "trend_training_hours_per_employee",
            "kpi_summary",
        ],
    )

    for label in chart_order:
        img_b64 = charts.get(label)
        if img_b64:
            p = doc.add_paragraph(label.replace("_", " ").title())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            img_buf = _decode_base64_image(img_b64)
            doc.add_picture(img_buf, width=Inches(4.0))

            doc.add_paragraph("")

    doc.add_paragraph("")

    # Narrative sections
    narrative_config = layout.get("narrative", {})
    sections_config = narrative_config.get("sections", [])

    # Create sections mapping
    sections_data = {
        "executive_summary": executive_summary,
        "workforce_composition_and_diversity": workforce_composition_and_diversity,
        "working_conditions_and_equal_opportunity": working_conditions_and_equal_opportunity,
        "training_and_development": training_and_development,
        "turnover_and_retention": turnover_and_retention,
        "health_and_safety": health_and_safety,
        "outlook_and_next_steps": outlook_and_next_steps,
    }

    for section_config in sections_config:
        title = section_config.get("title", "")
        key = section_config.get("key", "")
        content = sections_data.get(key, "")

        if title and content:
            if title == "Executive Summary":
                doc.add_heading(title, level=1)
            else:
                doc.add_heading(title, level=3)

            for para in content.split("\n\n"):
                doc.add_paragraph(para)

    # Closing
    closing_config = layout.get("closing", {})
    doc.add_heading(closing_config.get("title", "Closing"), level=1)
    doc.add_paragraph(closing)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return filename, base64.b64encode(buf.read()).decode("utf-8")


def generate_pdf_report(
    *,
    company_id: int,
    year: int,
    company_name: Optional[str],
    kpi_data: Dict[str, Any],
    charts: Dict[str, Optional[str]],
    executive_summary: str,
    workforce_composition_and_diversity: str,
    working_conditions_and_equal_opportunity: str,
    training_and_development: str,
    turnover_and_retention: str,
    health_and_safety: str,
    outlook_and_next_steps: str,
    closing: str,
    layout_template: Optional[Dict[str, Any]] = None,
) -> Tuple[str, str]:

    layout = get_layout("default") if layout_template is None else layout_template

    filename = f"S1_Report_{company_id}_{year}.pdf"
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36
    )
    styles = getSampleStyleSheet()
    story = []

    # Cover page
    cover_config = layout.get("cover_page", {})
    story.append(Spacer(1, 2 * inch))
    story.append(
        Paragraph(
            cover_config.get("title", "ESRS S1 Management Report"), styles["Title"]
        )
    )
    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph(f"{company_name or company_id}", styles["Heading2"]))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(f"Reporting Year: {year}", styles["Normal"]))
    story.append(Spacer(1, 0.5 * inch))

    # Logo
    if cover_config.get("show_logo", True):
        img = Image.new("RGB", (1024, 512), color=(230, 230, 230))
        img_buf = io.BytesIO()
        img.save(img_buf, format="PNG")
        img_buf.seek(0)
        story.append(RLImage(img_buf, width=5.5 * inch, height=3 * inch))
        story.append(Spacer(1, 0.7 * inch))

    # Date
    if cover_config.get("show_date", True):
        story.append(Paragraph(datetime.now().strftime("%B %d, %Y"), styles["Normal"]))
        story.append(Spacer(1, 2 * inch))

    # Confidential notice
    if cover_config.get("show_confidential", True):
        story.append(Paragraph("Confidential", styles["Italic"]))

    story.append(PageBreak())

    # KPI Summary table
    kpi_config = layout.get("kpi_summary", {})
    story.append(
        Paragraph(
            kpi_config.get("title", "KPI Summary"),
            styles["Heading1"],
        )
    )
    rows = _build_kpi_summary_rows(kpi_data)
    table = Table(rows, colWidths=[3.0 * inch, 3.0 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.2 * inch))

    # KPI Visualizations
    visuals_config = layout.get("visuals", {})
    story.append(
        Paragraph(
            visuals_config.get("title", "KPI Visualizations"),
            styles["Heading1"],
        )
    )

    # Display charts in configured order
    chart_order = visuals_config.get(
        "chart_order",
        [
            "workforce_by_gender",
            "training_hours_by_gender",
            "trend_training_hours_per_employee",
            "kpi_summary",
        ],
    )

    for label in chart_order:
        img_b64 = charts.get(label)
        if img_b64:
            story.append(Paragraph(label.replace("_", " ").title(), styles["Heading3"]))

            img_buf = _decode_base64_image(img_b64)
            story.append(RLImage(img_buf, width=4.5 * inch, height=2.6 * inch))

            story.append(Spacer(1, 0.25 * inch))

    story.append(Spacer(1, 0.2 * inch))

    story.append(PageBreak())

    # Narrative sections
    narrative_config = layout.get("narrative", {})
    sections_config = narrative_config.get("sections", [])

    # Create sections mapping
    sections_data = {
        "executive_summary": executive_summary,
        "workforce_composition_and_diversity": workforce_composition_and_diversity,
        "working_conditions_and_equal_opportunity": working_conditions_and_equal_opportunity,
        "training_and_development": training_and_development,
        "turnover_and_retention": turnover_and_retention,
        "health_and_safety": health_and_safety,
        "outlook_and_next_steps": outlook_and_next_steps,
    }

    for section_config in sections_config:
        title = section_config.get("title", "")
        key = section_config.get("key", "")
        content = sections_data.get(key, "")

        if title and content:
            if title == "Executive Summary":
                story.append(Paragraph(title, styles["Heading1"]))
            else:
                story.append(Paragraph(title, styles["Heading3"]))

            for para in content.split("\n\n"):
                story.append(Paragraph(para, styles["Normal"]))
                story.append(Spacer(1, 0.1 * inch))

    # Closing paragraph
    closing_config = layout.get("closing", {})
    story.append(Paragraph(closing_config.get("title", "Closing"), styles["Heading1"]))
    story.append(Paragraph(closing, styles["Normal"]))

    doc.build(story)
    buf.seek(0)
    return filename, base64.b64encode(buf.read()).decode("utf-8")
