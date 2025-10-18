import base64
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader
from app.file_export import generate_docx_report, generate_pdf_report


class TestESRSTagsAndVisuals:

    def test_esrs_s1_tags_in_docx(
        self, sample_kpi_data, sample_sections, sample_charts
    ):

        filename, content_b64 = generate_docx_report(
            company_id=1,
            year=2024,
            company_name="Test Company",
            kpi_data=sample_kpi_data,
            charts=sample_charts,
            executive_summary=sample_sections["executive_summary"],
            workforce_composition_and_diversity=sample_sections[
                "workforce_composition_and_diversity"
            ],
            working_conditions_and_equal_opportunity=sample_sections[
                "working_conditions_and_equal_opportunity"
            ],
            training_and_development=sample_sections["training_and_development"],
            turnover_and_retention=sample_sections["turnover_and_retention"],
            health_and_safety=sample_sections["health_and_safety"],
            outlook_and_next_steps=sample_sections["outlook_and_next_steps"],
            closing=sample_sections["closing"],
        )

        docx_content = base64.b64decode(content_b64)
        doc = Document(BytesIO(docx_content))
        text_content = "\n".join([p.text for p in doc.paragraphs])

        # Check for ESRS S1 section headers
        esrs_sections = [
            "Workforce Composition and Diversity",
            "Working Conditions and Equal Opportunity",
            "Training and Development",
            "Turnover and Retention",
            "Health and Safety",
            "Outlook and Next Steps",
        ]

        for section in esrs_sections:
            assert section in text_content, f"Missing ESRS S1 section: {section}"

        # Check for KPI summary table
        assert "KPI Summary" in text_content

    def test_visual_elements_in_docx(
        self, sample_kpi_data, sample_sections, sample_charts
    ):

        filename, content_b64 = generate_docx_report(
            company_id=1,
            year=2024,
            company_name="Test Company",
            kpi_data=sample_kpi_data,
            charts=sample_charts,
            executive_summary=sample_sections["executive_summary"],
            workforce_composition_and_diversity=sample_sections[
                "workforce_composition_and_diversity"
            ],
            working_conditions_and_equal_opportunity=sample_sections[
                "working_conditions_and_equal_opportunity"
            ],
            training_and_development=sample_sections["training_and_development"],
            turnover_and_retention=sample_sections["turnover_and_retention"],
            health_and_safety=sample_sections["health_and_safety"],
            outlook_and_next_steps=sample_sections["outlook_and_next_steps"],
            closing=sample_sections["closing"],
        )

        docx_content = base64.b64decode(content_b64)
        doc = Document(BytesIO(docx_content))
        text_content = "\n".join([p.text for p in doc.paragraphs])

        # Check for visualizations section
        assert "KPI Visualizations" in text_content

        if sample_charts.get("trend_training_hours_per_employee"):
            assert "Trend Training Hours Per Employee" in text_content

        # Check for chart titles (if charts are present)
        if sample_charts:
            chart_titles = [
                "Workforce By Gender",
                "Training Hours By Gender",
                "Trend Training Hours Per Employee",
            ]

            for title in chart_titles:
                if any(
                    key in sample_charts
                    for key in [
                        "workforce_by_gender",
                        "training_hours_by_gender",
                        "trend_training_hours_per_employee",
                    ]
                ):
                    assert title in text_content, f"Missing chart title: {title}"

    def test_esrs_s1_tags_in_pdf(
        self, sample_kpi_data, sample_sections, sample_charts
    ):
    
        filename, content_b64 = generate_pdf_report(
            company_id=1,
            year=2024,
            company_name="Test Company",
            kpi_data=sample_kpi_data,
            charts=sample_charts,
            executive_summary=sample_sections["executive_summary"],
            workforce_composition_and_diversity=sample_sections[
                "workforce_composition_and_diversity"
            ],
            working_conditions_and_equal_opportunity=sample_sections[
                "working_conditions_and_equal_opportunity"
            ],
            training_and_development=sample_sections["training_and_development"],
            turnover_and_retention=sample_sections["turnover_and_retention"],
            health_and_safety=sample_sections["health_and_safety"],
            outlook_and_next_steps=sample_sections["outlook_and_next_steps"],
            closing=sample_sections["closing"],
        )

        pdf_content = base64.b64decode(content_b64)
        pdf_reader = PdfReader(BytesIO(pdf_content))

        # Extract text from all pages
        full_text = ""
        for page in pdf_reader.pages:
            full_text += page.extract_text()

        # Check for ESRS S1 section headers
        esrs_sections = [
            "Workforce Composition and Diversity",
            "Working Conditions and Equal Opportunity",
            "Training and Development",
            "Turnover and Retention",
            "Health and Safety",
            "Outlook and Next Steps",
        ]

        for section in esrs_sections:
            assert section in full_text, f"Missing ESRS S1 section in PDF: {section}"

        # Check for KPI summary
        assert "KPI Summary" in full_text
