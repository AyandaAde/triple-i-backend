"""
Tests the complete flow from the /report endpoint through to
final document generation.
"""

import base64
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader
import pytest
from app.routers.report import create_report
from app.routers.report import ReportRequest


class TestFullIntegration:

    @pytest.mark.asyncio
    async def test_report_endpoint_pdf_integration(self, sample_kpi_data):
        
        # Create request payload for PDF
        request_payload = ReportRequest(
            company_id=1,
            year=2024,
            company_name="Integration Test Company",
            kpi_data=sample_kpi_data,
            historical_kpi_data=None,
            type="pdf",
        )

        # Test PDF generation
        result = await create_report(request_payload)

        # Validate response structure
        assert "sections" in result
        assert "charts" in result
        assert "file" in result

        # Validate sections
        sections = result["sections"]
        required_sections = [
            "executive_summary",
            "workforce_composition_and_diversity",
            "working_conditions_and_equal_opportunity",
            "training_and_development",
            "turnover_and_retention",
            "health_and_safety",
            "outlook_and_next_steps",
            "closing",
        ]

        for section in required_sections:
            assert section in sections, f"Missing section: {section}"
            assert sections[section], f"Empty section: {section}"

        # Validate file generation
        file_info = result["file"]
        assert file_info["name"] == "S1_Report_1_2024.pdf"
        assert file_info["base64"] is not None
        assert len(file_info["base64"]) > 0

        # Validate PDF content
        pdf_content = base64.b64decode(file_info["base64"])
        pdf_reader = PdfReader(BytesIO(pdf_content))
        assert len(pdf_reader.pages) > 0, "Generated PDF has no pages"

        # Check for key content
        first_page_text = pdf_reader.pages[0].extract_text()
        assert "Integration Test Company" in first_page_text
        assert "2024" in first_page_text

    @pytest.mark.asyncio
    async def test_report_endpoint_docx_integration(self, sample_kpi_data):

        # Create request payload for DOCX
        request_payload = ReportRequest(
            company_id=1,
            year=2024,
            company_name="Integration Test Company",
            kpi_data=sample_kpi_data,
            historical_kpi_data=None,
            type="docx",
        )

        result = await create_report(request_payload)

        # Validate file generation
        file_info = result["file"]
        assert file_info["name"] == "S1_Report_1_2024.docx"
        assert file_info["base64"] is not None
        assert len(file_info["base64"]) > 0

        # Validate DOCX content
        docx_content = base64.b64decode(file_info["base64"])
        doc = Document(BytesIO(docx_content))
        assert len(doc.paragraphs) > 0, "Generated DOCX is empty"

        # Check for key content
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Integration Test Company" in text_content
        assert "2024" in text_content
