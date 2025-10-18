"""
Tests that PDF and DOCX files are generated correctly and are non-empty

"""

import base64
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader
from app.file_export import generate_docx_report, generate_pdf_report


class TestFileGeneration:

    def test_docx_generation_valid_and_non_empty(
        self, sample_kpi_data, sample_sections, sample_charts
    ):

        filename, content_b64 = generate_docx_report(
            company_id=1,
            year=2024,
            company_name="Test Company",
            kpi_data=sample_kpi_data,
            charts=sample_charts,
            executive_summary=sample_sections["executive_summary"],
            workforce_composition_and_diversity=sample_sections[
                "workforce_composition_and_diversity"
            ],
            working_conditions_and_equal_opportunity=sample_sections[
                "working_conditions_and_equal_opportunity"
            ],
            training_and_development=sample_sections["training_and_development"],
            turnover_and_retention=sample_sections["turnover_and_retention"],
            health_and_safety=sample_sections["health_and_safety"],
            outlook_and_next_steps=sample_sections["outlook_and_next_steps"],
            closing=sample_sections["closing"],
        )

        # Validate filename
        assert filename == "S1_Report_1_2024.docx"
        assert content_b64 is not None
        assert len(content_b64) > 0

        # Validate DOCX content
        docx_content = base64.b64decode(content_b64)
        doc = Document(BytesIO(docx_content))

        # Check document has content
        assert len(doc.paragraphs) > 0, "DOCX document is empty"

        # Check for key elements
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Test Company" in text_content
        assert "2024" in text_content

    def test_pdf_generation_valid_and_non_empty(
        self, sample_kpi_data, sample_sections, sample_charts
    ):
    
        filename, content_b64 = generate_pdf_report(
            company_id=1,
            year=2024,
            company_name="Test Company",
            kpi_data=sample_kpi_data,
            charts=sample_charts,
            executive_summary=sample_sections["executive_summary"],
            workforce_composition_and_diversity=sample_sections[
                "workforce_composition_and_diversity"
            ],
            working_conditions_and_equal_opportunity=sample_sections[
                "working_conditions_and_equal_opportunity"
            ],
            training_and_development=sample_sections["training_and_development"],
            turnover_and_retention=sample_sections["turnover_and_retention"],
            health_and_safety=sample_sections["health_and_safety"],
            outlook_and_next_steps=sample_sections["outlook_and_next_steps"],
            closing=sample_sections["closing"],
        )

        # Validate filename
        assert filename == "S1_Report_1_2024.pdf"
        assert content_b64 is not None
        assert len(content_b64) > 0

        # Validate PDF content
        pdf_content = base64.b64decode(content_b64)
        pdf_reader = PdfReader(BytesIO(pdf_content))

        # Check PDF has pages
        assert len(pdf_reader.pages) > 0, "PDF document has no pages"

        # Check for key content in first page
        first_page_text = pdf_reader.pages[0].extract_text()
        assert "Test Company" in first_page_text
        assert "2024" in first_page_text
